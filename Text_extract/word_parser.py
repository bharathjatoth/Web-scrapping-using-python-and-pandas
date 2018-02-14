from docx.shared import Inches
from docx.shared import Pt
import datetime
from docx import Document
from docx import *
from pathlib import Path
import docx
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
# from docx.xml.text.paragraph import
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent.tc
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
fullText = []  #for the para text
fullText1 = []  #for heading1 
fullText2 = []  #
fullText5 = []  #for heading1 + heading2
fullText6 = []  # only heading 2
def parse_word():
    pathlist = Path(Path_of_the_Foleder_with_word_doc).glob('**/*.docx')
    x1 = ''
    x2 = ''
    x3 = ''
    x4 = ''
    y1 = ' '
    x = ''
    y = y2 = ''
    i = 0
    #search for each word doc
    for path in pathlist:
        doc = Document(str(path))
        document = docx.Document(str(path))
        #get tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para2 in cell.paragraphs:
                        x1 = x1 + (para2.text) + '\t'
                x4 = x4 + x1 + '\n'
                x1 = ''
            fullText2.append(str(x4))
            x4 = ''
        #get the para text with respective headers and para text
        for para1 in doc.paragraphs:
            if para1.style.name.startswith('Heading 1'):
                y2 = str(para1.text) + '\n'
                fullText1.append(y2)
                # print(y2)
            if para1.style.name.startswith('Heading 2'):
                y = y2 + str(para1.text) + '\n'
                fullText5.append(y)
                fullText6.append(para1.text)
                # fullText[len(fullText) - 1] = fullText[len(fullText) - 1] + str(y)
                # print(y)
                y = ''
            if para1.style.name.startswith('Normal'):
                y1 = str(para1.text)
                # print(para1.alignment)
                # print(y1)
                # print(len(fullText))
                fullText[len(fullText) - 1] = fullText[len(fullText) - 1] + str(y1)
                y1 = ''
    return fullText,fullText2,fullText5,fullText6
fullText3 = []   #the output text which displays the text in format of heading1->heading2->para.text(if there is table we display that)
fullText4 = []   #stores the headings of the current text we need to do it in the upcoming sessions where we develop a AI based application 
def parse_table():
    parse_word()
    j = 0
    m = 0
    pathlist = Path(Path_of_the_Foleder_with_word_doc).glob('**/*.docx')
    for path in pathlist:
        document = docx.Document(str(path))
        i = 0
        y2 = ''
        k = 0
        for para1 in document.paragraphs:
            if para1.style.name.startswith('Heading 1'):
                y2 = str(para1.text)
        for block in iter_block_items(document):
            i = i + 1
            y3 = ''
            #see whether the para is a instance of para or header or table and map to the respective text collected in the past i.e in fulltext lists 
            if isinstance(block,Paragraph):
                if len(str(block.text)) > 1:
                    y3 = str(fullText[:-1])
                    if str(block.text) in fullText6[m:]:
                        m = m+1
                        fullText3.append(str(block.text))
                        fullText4.append(str(block.text))
                    else:
                        if str(block.text) == y2:
                            print(block.text)
                        else:
                            fullText4.append(str(fullText6[m-1]) + '\n')
                            fullText3.append(str(fullText6[m-1]) + '\n' + block.text)
            if isinstance(block,Table):
                fullText4.append(str(fullText6[m-1]) + '\n')
                fullText3.append(str(fullText6[m-1]) + '\n' + str(fullText2[j]))
                j = j+1
    return fullText3,fullText4
if __name__ == '__main__':
    parse_table()
for i in range(len(fullText3)):
    print(i)
    print(fullText3[i])
    print(fullText4[i])
