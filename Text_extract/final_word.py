from docx.shared import Inches
from docx.shared import Pt
import datetime
from docx import Document
from docx import *
from pathlib import Path
import docx
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
# from docx.xml.text.paragraph import
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
doc = docx.Document(r'Document to parse')
def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent.tc
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
fullText = []  #for the para text
fullText1 = []  #for heading1
fullText2 = []  #
fullText5 = []  #for heading1 + heading2
fullText6 = []  # only heading 2
ques = []
ans = []
final_ques = []
final_ans = []
fullText3 = []   #the output text which displays the text in format of heading1->heading2->para.text(if there is table we display that)
fullText4 = []   #stores the headings of the current text we need to do it in the upcoming sessions where we develop a AI based application
head = []
tail = []
def parse_table(y2,question_col,answer_col,skip_rows,comment_col):
    table1 = False
    # a5 = parse_word(question_col,answer_col)
    # print("after")
    # print(a5)
    j = m = i = k = 0
    document = docx.Document(str(r'C:\Users\jatoth.kumar\Desktop\Hotscan - Copy\Hotscan - Copy\CGI_Response_to_AIB_SanctionScreening__RFI_20140910 V0 2 je - GOOD.docx'))
    for para1 in document.paragraphs:
        if para1.style.name.startswith('Heading 1'):
            y5 = str(para1.text)
    list3 = []
    list3 = document.tables
    print("printing list3")
    print(list3)
    print(len(list3))
    for block in iter_block_items(document):
        i = i + 1
        y3 = ''
        if isinstance(block,Paragraph):
            if len(str(block.text)) > 1:
                if str(block.text) == y2:
                    table1 = True
        if isinstance(block,Table):
            print("in inner table loop")
            row1 = 0
            if table1:
                if block in doc.tables:
                    print(collect[list3.index(block)])
                ques_col = []
                ans_col = []
                for table2 in block.rows:
                    for cell in table2.cells:
                        x1 = (cell.text) + '\t'
                        if row1 == question_col:
                            quest = cell.text
                            ques_col.append(quest)
                            print(quest)
                            print('question')
                        if row1 == answer_col:
                            answ = cell.text
                            ans_col.append(answ)
                            print(answ)
                            print("answer")
                        row1 = row1+1
                    row1 = 0
            print("table")
            table1 = False
    x2 = collect[3]
    print(len(ans_col),len(ques_col))
    quest, answer = x2['question'], x2['answer']
    quest, answer = quest[skip_rows:], answer[skip_rows:]
    # print(len(x2['question']), len(x2['answer']))
    # print(len(quest), len(answer))
    # for i in range(len(quest)):
    #     print(i)
    #     print(quest[i])
if __name__ == '__main__':
    y1 = 'Statement of Requirements '
    question_col = 1
    answer_col = 2
    skip_rows = 2
    parse_table(y1,question_col,answer_col,skip_rows,comment_col)
