from docx.shared import Inches
from docx.shared import Pt
import datetime
from docx import Document
from docx import *
from pathlib import Path
import docx
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
# from docx.xml.text.paragraph import
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
doc = docx.Document(docx location)
collect = {}
def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent.tc
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
fullText = []  #for the para text
fullText1 = []  #for heading1
fullText2 = []  #
fullText5 = []  #for heading1 + heading2
fullText6 = []  # only heading 2
ques = []
ans = []
final_ques = []
final_ans = []
def parse_word():
    global collect
    x1 = x2 = x3 = x4 = y1 = x = y = y2 = ''
    #search for each word doc
    row1 = k1 = 0
    doc = Document(str(word location))
    document = docx.Document(str(word location))
    #get tables
    quest = ''
    answ = ''
    print(len(doc.tables))
    collect = dict([i,0] for i in range(len(doc.tables)))
    for table in doc.tables:
        ques1 = []
        ans1 = []
        for row in table.rows:
            for cell in row.cells:
                x1 = (cell.text) + '\t'
                if row1 == 1:
                    quest = cell.text
                    # print(quest)
                    # print('----')
                if row1 == 2:
                    answ = cell.text
                    # print(answ)
                row1 = row1+1
            ans1.append(answ)
            ques1.append(quest)
            row1 = 0
            ques = answ = ''
            x4 = x4 + x1 + '\n'
            x1 = ''
        fullText2.append(ans1)
        x4 = ''
        print(k1)
        collect[k1] = {"question":ques1,"answer":ans1}
        print(len(ans1), len(ques1))
        k1 = k1+1
    #get the para text with respective headers and para text
    for para1 in doc.paragraphs:
        if para1.style.name.startswith('Heading 1'):
            y2 = str(para1.text) + '\n'
            fullText1.append(para1.text)
            # print(y2)
        if para1.style.name.startswith('Heading 2'):
            y = y2 + str(para1.text) + '\n'
            fullText5.append(y)
            fullText6.append(para1.text)
            # fullText[len(fullText) - 1] = fullText[len(fullText) - 1] + str(y)
            # print(y)
            y = ''
        if para1.style.name.startswith('Normal'):
            y1 = str(para1.text)
            # print(para1.alignment)
            # print(y1)
            # print(len(fullText))
            fullText.append(str(y1))
            y1 = ''
    print("in the parw")
    print(collect)
    return collect
fullText3 = []   #the output text which displays the text in format of heading1->heading2->para.text(if there is table we display that)
fullText4 = []   #stores the headings of the current text we need to do it in the upcoming sessions where we develop a AI based application
head = []
tail = []
def parse_table():
    table1 = False
    a5 = parse_word()
    print("after")
    print(a5)
    j = 0
    m = 0
    document = docx.Document(docx loaction)
    i = 0
    y2 = ''
    k = 0
    inner = False
    y2 = #name of the header under which we have table to extract
    for para1 in document.paragraphs:
        if para1.style.name.startswith('Heading 1'):
            y5 = str(para1.text)
    for block in iter_block_items(document):
        i = i + 1
        y3 = ''
        # print(block)
        #see whether the para is a instance of para or header or table and map to the respective text collected in the past i.e in fulltext lists
        if isinstance(block,Paragraph):
            if len(str(block.text)) > 1:
                if str(block.text) == y2:
                    inner = True
                    table1 = True
                    tail.append(fullText1.index(block.text))
        if isinstance(block,Table):
            if table1:
                print(tail[0])
                print(collect[tail[0]])
                print("table")
            table1 = False
    return collect
if __name__ == '__main__':
    parse_table()
x2 = collect[3]
print(len(x2['question']),len(x2['answer']))
